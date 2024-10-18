from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openpyxl import load_workbook
import pandas as pd


# Функция для изменения шрифта и размера шрифта
def set_font(run, font_name, font_size, italic=False, bold=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.italic = italic
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


# Функция для установки отступов и междустрочного интервала
def set_paragraph_format(paragraph, left_indent=0, right_indent=0, first_line_indent=1.25, line_spacing=22,
                         space_after=0, space_before=0):
    paragraph_format = paragraph.paragraph_format
    paragraph_format.left_indent = Cm(left_indent)
    paragraph_format.right_indent = Cm(right_indent)
    paragraph_format.first_line_indent = Cm(first_line_indent)
    paragraph_format.line_spacing = Pt(line_spacing)
    paragraph_format.space_after = Cm(space_after)
    paragraph_format.space_before = Cm(space_before)


# Функция для добавления заголовка
def add_header(doc, header_text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(header_text)
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY


# Функция для добавления таблицы с учётом объединения ячеек
def add_table(doc, df, merged_ranges):
    table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
    table.style = 'Table Grid'

    # Добавление заголовков таблицы
    hdr_cells = table.rows[0].cells
    for i, column_name in enumerate(df.columns):
        cell_paragraph = hdr_cells[i].paragraphs[0]
        cell_paragraph.text = column_name
        cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        for run in cell_paragraph.runs:
            set_font(run, 'Times New Roman', 12)
        set_paragraph_format(cell_paragraph, left_indent=0.0, right_indent=0.0, first_line_indent=0.0,
                             line_spacing=18, space_after=0, space_before=0)

    # Добавление строк таблицы
    for index, row in df.iterrows():
        row_cells = table.rows[index + 1].cells
        for i, value in enumerate(row):
            # Проверяем, что значение не является пустым
            if pd.notna(value) and str(value).strip() != "":
                cell_paragraph = row_cells[i].paragraphs[0]
                cell_paragraph.text = str(value)
                cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                for run in cell_paragraph.runs:
                    set_font(run, 'Times New Roman', 12)
                set_paragraph_format(cell_paragraph, left_indent=0.0, right_indent=0.0, first_line_indent=0.0,
                                     line_spacing=18, space_after=0, space_before=0)

    # Объединение ячеек в Word на основе объединённых диапазонов из Excel
    for merged_range in merged_ranges:
        min_col, min_row, max_col, max_row = merged_range.bounds
        for row in range(min_row, max_row + 1):
            for col in range(min_col, max_col + 1):
                # Объединяем ячейки в Word
                if row == min_row and col == min_col:
                    start_cell = table.cell(min_row - 1, min_col - 1)
                    end_cell = table.cell(max_row - 1, max_col - 1)
                    start_cell.merge(end_cell)


# Функция для вставки разрыва страницы
def insert_page_break(doc):
    doc.add_page_break()


# Чтение Excel с учетом объединённых ячеек
def read_excel_with_merged_cells(filename, sheet_name):
    # Загружаем Excel-файл с помощью openpyxl
    wb = load_workbook(filename, data_only=True)
    ws = wb[sheet_name]

    # Сохраняем данные в список строк (list of lists)
    data = []
    for row in ws.iter_rows(values_only=True):
        data.append(list(row))

    # Получаем объединенные ячейки
    merged_ranges = ws.merged_cells.ranges

    # Проходим по каждому объединенному диапазону и заполняем его данные
    for merged_range in merged_ranges:
        # Получаем диапазон объединённых ячеек
        min_col, min_row, max_col, max_row = merged_range.bounds

        # Получаем значение из первой ячейки диапазона
        merged_value = ws.cell(row=min_row, column=min_col).value

        # Присваиваем это значение только первой ячейке, остальные оставляем пустыми
        for row in range(min_row - 1, max_row):
            for col in range(min_col - 1, max_col):
                if row == min_row - 1 and col == min_col - 1:
                    data[row][col] = merged_value
                else:
                    data[row][col] = ''  # Оставляем остальные ячейки пустыми

    # Преобразуем данные в DataFrame pandas
    df = pd.DataFrame(data[1:], columns=data[0])
    return df, merged_ranges


# Основная программа для создания документа Word
def create_document():
    # Создаем новый документ
    doc = Document()

    # Настраиваем поля страницы
    section = doc.sections[0]
    section.left_margin = Cm(2)
    section.right_margin = Cm(1)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # Добавляем текст
    text = ['ОПРОСНЫЙ ЛИСТ №001',
            'для заказа колонных аппаратов',
            'Условное обозначение аппарата C-321 – вертикальный колонный аппарат Количество – 1']

    for line in text:
        paragraph = doc.add_paragraph(line)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in paragraph.runs:
            set_font(run, 'Times New Roman', 12)
        set_paragraph_format(paragraph, left_indent=0.0, right_indent=0.0, first_line_indent=1.25, line_spacing=22,
                             space_after=0, space_before=0)

    # Чтение данных из Excel с учетом объединённых ячеек
    df1, merged_ranges = read_excel_with_merged_cells('database/device/1.xlsx', '1')

    # Добавление заголовка и таблицы
    add_header(doc, '1. Климатические условия в районе строительства')
    add_table(doc, df1, merged_ranges)

    # Сохранение документа
    doc.save('ОЛ #1.docx')


# Запуск программы
create_document()
