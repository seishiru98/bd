from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION, WD_ORIENT

import pandas as pd
from openpyxl import load_workbook

import datetime

# -----------------------------------------------------------------------------------------------------------------------

doc = Document()

section = doc.sections[0]

section.left_margin = Cm(2)  # Левое поле
section.right_margin = Cm(1)  # Правое поле
section.top_margin = Cm(2)  # Верхнее поле
section.bottom_margin = Cm(2)  # Нижнее поле


# -----------------------------------------------------------------------------------------------------------------------
# Функции

def set_font(run, font_name, font_size, italic=False, bold=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.italic = italic
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    # Это нужно для корректного отображения шрифта на всех платформах
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.append(rFonts)


def set_paragraph_format(paragraph, left_indent=0, right_indent=0, first_line_indent=1.25, line_spacing=22,
                         space_after=0, space_before=0):
    paragraph_format = paragraph.paragraph_format
    paragraph_format.left_indent = Cm(left_indent)
    paragraph_format.right_indent = Cm(right_indent)
    paragraph_format.first_line_indent = Cm(first_line_indent)
    paragraph_format.line_spacing = Pt(line_spacing)
    paragraph_format.space_after = Cm(space_after)
    paragraph_format.space_before = Cm(space_before)


def add_header(doc, header_text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(header_text)
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY


def add_table(doc, df, merged_ranges):
    # Общая ширина таблицы в сантиметрах
    total_width = Cm(18.5)  # Примерная ширина текста на странице A4 с полями

    # Определяем таблицу и стиль
    table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
    table.style = 'Table Grid'

    # Игнорируем первую строку, чтобы определить максимальную длину текста в каждой колонке
    column_widths = [0] * len(df.columns)

    # Добавляем строки с данными и определяем максимальную длину текста в каждом столбце
    for index, row in df.iterrows():
        for i, value in enumerate(row):
            if pd.notna(value) and str(value).strip() != "":
                column_widths[i] = max(column_widths[i], len(str(value)))

    # Вычисляем общую длину текста для пропорциональной настройки ширины столбцов
    total_text_length = sum(column_widths)

    # Добавление заголовков таблицы
    hdr_cells = table.rows[0].cells
    for i, column_name in enumerate(df.columns):
        hdr_cells[i].width = Cm(
            total_width.cm * (column_widths[i] / total_text_length))  # Задаем ширину на основе данных
        cell_paragraph = hdr_cells[i].paragraphs[0]
        cell_paragraph.text = column_name
        cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        for run in cell_paragraph.runs:
            set_font(run, 'Times New Roman', 8)  # Шрифт для заголовков
        set_paragraph_format(cell_paragraph, left_indent=0.0, right_indent=0.0, first_line_indent=0.0,
                             line_spacing=11.5, space_after=0, space_before=0)

    # Добавление строк таблицы
    for index, row in df.iterrows():
        row_cells = table.rows[index + 1].cells
        for i, value in enumerate(row):
            row_cells[i].width = Cm(
                total_width.cm * (column_widths[i] / total_text_length))  # Задаем ширину для ячеек на основе данных
            if pd.notna(value) and str(value).strip() != "":
                cell_paragraph = row_cells[i].paragraphs[0]
                cell_paragraph.text = str(value)
                cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                for run in cell_paragraph.runs:
                    set_font(run, 'Times New Roman', 8)
                set_paragraph_format(cell_paragraph, left_indent=0.0, right_indent=0.0, first_line_indent=0.0,
                                     line_spacing=11.5, space_after=0, space_before=0)

    # Объединение ячеек в Word на основе объединённых диапазонов из Excel
    for merged_range in merged_ranges:
        min_col, min_row, max_col, max_row = merged_range.bounds
        for row in range(min_row, max_row + 1):
            for col in range(min_col, max_col + 1):
                if row == min_row and col == min_col:
                    start_cell = table.cell(min_row - 1, min_col - 1)
                    end_cell = table.cell(max_row - 1, max_col - 1)
                    start_cell.merge(end_cell)


def insert_page_break(doc):
    doc.add_page_break()


def read_excel_with_merged_cells(filename, sheet_name):
    # Загружаем Excel-файл с помощью openpyxl
    wb = load_workbook(filename, data_only=True)
    ws = wb[sheet_name]

    # Сохраняем данные в список строк (list of lists)
    data = []
    for row in ws.iter_rows(values_only=True):
        data.append(list(row))

    # Получаем объединенные ячейки
    merged_ranges = ws.merged_cells.ranges

    # Проходим по каждому объединенному диапазону и заполняем его данные
    for merged_range in merged_ranges:
        # Получаем диапазон объединённых ячеек
        min_col, min_row, max_col, max_row = merged_range.bounds

        # Получаем значение из первой ячейки диапазона
        merged_value = ws.cell(row=min_row, column=min_col).value

        # Присваиваем это значение только первой ячейке, остальные оставляем пустыми
        for row in range(min_row - 1, max_row):
            for col in range(min_col - 1, max_col):
                if row == min_row - 1 and col == min_col - 1:
                    data[row][col] = merged_value
                else:
                    data[row][col] = ''  # Оставляем остальные ячейки пустыми

    # Преобразуем данные в DataFrame pandas
    df = pd.DataFrame(data[1:], columns=data[0])
    return df, merged_ranges


class Counter:
    def __init__(self, start_value, step):
        self.value = start_value
        self.step = step

    def increment(self):
        current_value = self.value
        self.value += self.step
        return current_value


class HeadingCounter(Counter):
    def __init__(self, start_value, paragraph_counter, table_counter, fig_counter):
        super().__init__(start_value, 1)
        self.paragraph_counter = paragraph_counter
        self.table_counter = table_counter
        self.fig_counter = fig_counter

    def increment(self):
        current_value = super().increment()
        self.paragraph_counter.reset(current_value + 0.1)
        self.table_counter.reset(current_value + 0.1)
        self.fig_counter.reset(current_value + 0.1)
        return current_value


class ParagraphCounter(Counter):
    def reset(self, new_start_value):
        self.value = new_start_value


class TableCounter(Counter):
    def reset(self, new_start_value):
        self.value = new_start_value


class FigCounter(Counter):
    def reset(self, new_start_value):
        self.value = new_start_value


# Инициализация счетчиков
n_heading = 1
n_paragraph_start = n_heading + 0.1
n_table_start = n_heading + 0.1
n_fig_start = n_heading + 0.1

par_counter = ParagraphCounter(n_paragraph_start, 0.1)
table_counter = TableCounter(n_table_start, 0.1)
fig_counter = FigCounter(n_fig_start, 0.1)
head_counter = HeadingCounter(n_heading, par_counter, table_counter, fig_counter)


def read_excel_data(filename, sheet_name):
    try:
        return pd.read_excel(filename, sheet_name=sheet_name)
    except FileNotFoundError:
        print(f"ОШИБКА: Файл {filename} не найден.")
    except ValueError:
        print(f"ОШИБКА: Лист {sheet_name} не найден в файле {filename}.")
    except Exception as e:
        print(f"ОШИБКА: Произошла ошибка при чтении файла {filename}: {e}")


# -----------------------------------------------------------------------------------------------------------------------
ch_1 = head_counter.increment()

heading = doc.add_heading(f'{ch_1:.0f} ТАБЛИЦЫ ПОТОКОВ ', level=1)
heading.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

for run in heading.runs:
    set_font(run, 'Times New Roman', 14)
    set_paragraph_format(heading, left_indent=0.0, right_indent=0.0, first_line_indent=1.25, line_spacing=22,
                         space_after=0, space_before=0)

text = [f''
       ]

for line in text:
    paragraph_after_break = doc.add_paragraph(line)
    paragraph_after_break.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph_after_break.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph_after_break, left_indent=0.0, right_indent=0.0, first_line_indent=1.25,
                         line_spacing=22, space_after=0, space_before=0)

table1_1 = table_counter.increment()

df1, merged_ranges = read_excel_with_merged_cells('term.xlsx', '1')
add_header(doc, f'Таблица {table1_1:.1f} – Поток № 1 ')
add_table(doc, df1, merged_ranges)

text = [f''
       ]

for line in text:
    paragraph_after_break = doc.add_paragraph(line)
    paragraph_after_break.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph_after_break.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph_after_break, left_indent=0.0, right_indent=0.0, first_line_indent=1.25,
                         line_spacing=22, space_after=0, space_before=0)

table1_2 = table_counter.increment()

df2, merged_ranges = read_excel_with_merged_cells('term.xlsx', '2')
add_header(doc, f'Таблица {table1_2:.1f} – Поток № 2 ')
add_table(doc, df2, merged_ranges)

text = [f''
       ]

for line in text:
    paragraph_after_break = doc.add_paragraph(line)
    paragraph_after_break.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph_after_break.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph_after_break, left_indent=0.0, right_indent=0.0, first_line_indent=1.25,
                         line_spacing=22, space_after=0, space_before=0)
# -----------------------------------------------------------------------------------------------------------------------

# Сохраняем документ
doc.save('потоки.docx')
