
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION, WD_ORIENT

import pandas as pd
from openpyxl import load_workbook

import datetime
#-----------------------------------------------------------------------------------------------------------------------

doc = Document()

section = doc.sections[0]

section.left_margin = Cm(2)  # Левое поле
section.right_margin = Cm(1)  # Правое поле
section.top_margin = Cm(2)  # Верхнее поле
section.bottom_margin = Cm(2)  # Нижнее поле
#-----------------------------------------------------------------------------------------------------------------------
# Функции

def set_font(run, font_name, font_size, italic=False, bold=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.italic = italic
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    # Это нужно для корректного отображения шрифта на всех платформах
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.append(rFonts)


def set_paragraph_format(paragraph, left_indent=0, right_indent=0, first_line_indent=1.25, line_spacing=22,
                         space_after=0, space_before=0):
    paragraph_format = paragraph.paragraph_format
    paragraph_format.left_indent = Cm(left_indent)
    paragraph_format.right_indent = Cm(right_indent)
    paragraph_format.first_line_indent = Cm(first_line_indent)
    paragraph_format.line_spacing = Pt(line_spacing)
    paragraph_format.space_after = Cm(space_after)
    paragraph_format.space_before = Cm(space_before)


def add_header(doc, header_text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(header_text)
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY


def add_table(doc, df, start_row, end_row, merged_ranges, include_header=True):
    total_width = Cm(25.5)  # Общая ширина таблицы (например, вся ширина страницы)
    max_col_widths = []

    # Определим максимальную длину текста для каждой колонки
    for col in df.columns:
        max_width = max(df[col][start_row:end_row].apply(lambda x: len(str(x)) if x is not None else 0))
        max_col_widths.append(max_width)

    # Пропорциональная ширина колонок относительно их максимального содержания
    total_content_width = sum(max_col_widths)
    col_width_ratios = [width / total_content_width for width in max_col_widths]
    col_widths = [total_width * ratio for ratio in col_width_ratios]

    # Создаем таблицу с количеством столбцов, равным числу колонок в DataFrame
    num_columns = len(df.columns)
    table = doc.add_table(rows=0, cols=num_columns)
    table.style = 'Table Grid'

    # Добавляем заголовок таблицы, если включен параметр include_header
    if include_header:
        hdr_cells = table.add_row().cells
        for i, column_name in enumerate(df.columns):
            hdr_cells[i].width = col_widths[i]  # Применяем пропорциональную ширину к заголовкам
            cell_paragraph = hdr_cells[i].paragraphs[0]
            cell_paragraph.text = str(column_name)
            cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            for run in cell_paragraph.runs:
                set_font(run, 'Times New Roman', 12, bold=False)
            set_paragraph_format(cell_paragraph, left_indent=0.0, right_indent=0.0,
                                 first_line_indent=0.0, line_spacing=18, space_after=0, space_before=0)

    # Добавление строк таблицы
    for index in range(start_row, end_row):
        row = df.iloc[index]
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].width = col_widths[i]  # Применяем пропорциональную ширину к каждой ячейке строки
            cell_paragraph = row_cells[i].paragraphs[0]
            cell_paragraph.text = str(value)
            cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            for run in cell_paragraph.runs:
                set_font(run, 'Times New Roman', 12)
            set_paragraph_format(cell_paragraph, left_indent=0.0, right_indent=0.0,
                                 first_line_indent=0.0, line_spacing=18, space_after=0, space_before=0)

    # Корректировка для индексации строк
    header_offset = 1 if include_header else 0

    # Обработка объединённых ячеек
    for merged_range in merged_ranges:
        min_col, min_row, max_col, max_row = merged_range.bounds
        min_col -= 1  # Преобразование к нулевой базе
        max_col -= 1

        # Корректировка индексов строк для DataFrame
        min_row_df = min_row - 2  # -2, потому что DataFrame начинается с Excel строки 2 и индексируется с 0
        max_row_df = max_row - 2

        # Проверяем, попадает ли объединённый диапазон в текущий диапазон строк
        if min_row_df >= start_row and max_row_df < end_row:
            start_cell_row = (min_row_df - start_row) + header_offset
            end_cell_row = (max_row_df - start_row) + header_offset

            start_cell_col = min_col
            end_cell_col = max_col

            start_cell = table.cell(int(start_cell_row), int(start_cell_col))
            end_cell = table.cell(int(end_cell_row), int(end_cell_col))
            start_cell.merge(end_cell)

            # Удаляем лишние пустые параграфы из объединенной ячейки
            for paragraph in start_cell.paragraphs:
                if not paragraph.text.strip():
                    p = paragraph._element
                    p.getparent().remove(p)
                    p._p = p._element = None


def insert_page_break(doc):
    doc.add_page_break()


def read_excel_with_merged_cells(filename, sheet_name):
    # Загружаем Excel-файл с помощью openpyxl
    wb = load_workbook(filename, data_only=True)
    ws = wb[sheet_name]

    # Сохраняем данные в список строк (list of lists)
    data = []
    for row in ws.iter_rows(values_only=True):
        data.append(list(row))

    # Получаем объединенные ячейки
    merged_ranges = ws.merged_cells.ranges

    # Проходим по каждому объединенному диапазону и заполняем его данные
    for merged_range in merged_ranges:
        # Получаем диапазон объединённых ячеек
        min_col, min_row, max_col, max_row = merged_range.bounds

        # Получаем значение из первой ячейки диапазона
        merged_value = ws.cell(row=min_row, column=min_col).value

        # Присваиваем это значение только первой ячейке, остальные оставляем пустыми
        for row in range(min_row - 1, max_row):
            for col in range(min_col - 1, max_col):
                if row == min_row - 1 and col == min_col - 1:
                    data[row][col] = merged_value
                else:
                    data[row][col] = ''  # Оставляем остальные ячейки пустыми

    # Преобразуем данные в DataFrame pandas
    df = pd.DataFrame(data[1:], columns=data[0])
    return df, merged_ranges


class Counter:
    def __init__(self, start_value, step):
        self.value = start_value
        self.step = step

    def increment(self):
        current_value = self.value
        self.value += self.step
        return current_value

class HeadingCounter(Counter):
    def __init__(self, start_value, paragraph_counter, table_counter, fig_counter):
        super().__init__(start_value, 1)
        self.paragraph_counter = paragraph_counter
        self.table_counter = table_counter
        self.fig_counter = fig_counter

    def increment(self):
        current_value = super().increment()
        self.paragraph_counter.reset(current_value + 0.1)
        self.table_counter.reset(current_value + 0.1)
        self.fig_counter.reset(current_value + 0.1)
        return current_value

class ParagraphCounter(Counter):
    def reset(self, new_start_value):
        self.value = new_start_value

class TableCounter(Counter):
    def reset(self, new_start_value):
        self.value = new_start_value

class FigCounter(Counter):
    def reset(self, new_start_value):
        self.value = new_start_value

# Инициализация счетчиков
n_heading = 1
n_paragraph_start = n_heading + 0.1
n_table_start = n_heading + 0.1
n_fig_start = n_heading + 0.1

par_counter = ParagraphCounter(n_paragraph_start, 0.1)
table_counter = TableCounter(n_table_start, 0.1)
fig_counter = FigCounter(n_fig_start, 0.1)
head_counter = HeadingCounter(n_heading, par_counter, table_counter, fig_counter)

def read_excel_data(filename, sheet_name):
    try:
        return pd.read_excel(filename, sheet_name=sheet_name)
    except FileNotFoundError:
        print(f"ОШИБКА: Файл {filename} не найден.")
    except ValueError:
        print(f"ОШИБКА: Лист {sheet_name} не найден в файле {filename}.")
    except Exception as e:
        print(f"ОШИБКА: Произошла ошибка при чтении файла {filename}: {e}")


#-----------------------------------------------------------------------------------------------------------------------
# Переменные
database = pd.read_excel('database.xlsx', sheet_name='1')

work_time = database.iloc[0, 3]
print(work_time)

mass_frac_tiols = database.iloc[15, 4]
print(mass_frac_tiols)
mass_frac_sulphur = database.iloc[16, 4]
print(mass_frac_sulphur)

ppm_tiols = mass_frac_tiols * 10000
print(ppm_tiols)
ppm_sulphur = mass_frac_sulphur * 10000
print(ppm_sulphur)

min_flow_rate = database.iloc[18, 3]
print(min_flow_rate)
flow_rate = database.iloc[19, 3]
print(flow_rate)

MPS_calc_p, MPS_calc_t = database.iloc[55, 1], database.iloc[55, 2]
print(MPS_calc_p, MPS_calc_t)
MPS_work_p, MPS_work_t = database.iloc[56, 1], database.iloc[56, 2]
print(MPS_work_p, MPS_work_t)

LPS_calc_p, LPS_calc_t = database.iloc[60, 1], database.iloc[60, 2]
print(LPS_calc_p, LPS_calc_t)
LPS_work_p, LPS_work_t = database.iloc[61, 1], database.iloc[61, 2]
print(LPS_work_p, LPS_work_t)

water_direct_p, water_direct_t = database.iloc[65, 1], database.iloc[65, 2]
print(water_direct_p, water_direct_t)
water_reversed_p, water_reversed_t = database.iloc[75, 1], database.iloc[75, 2]
print(water_reversed_p, water_reversed_t)

LPG_Nitrogen_calc_p, LPG_Nitrogen_calc_t = database.iloc[105, 1], database.iloc[105, 2]
print(LPG_Nitrogen_calc_p, LPG_Nitrogen_calc_t)
LPG_Nitrogen_work_p, LPG_Nitrogen_work_t = database.iloc[106, 1], database.iloc[106, 2]
print(LPG_Nitrogen_work_p, LPG_Nitrogen_work_t)

HPG_Nitrogen_calc_p, HPG_Nitrogen_calc_t = database.iloc[112, 1], database.iloc[112, 2]
print(HPG_Nitrogen_calc_p, HPG_Nitrogen_calc_t)
HPG_Nitrogen_work_p, HPG_Nitrogen_work_t = database.iloc[113, 1], database.iloc[113, 2]
print(HPG_Nitrogen_work_p, HPG_Nitrogen_work_t)

air_calc_p, air_calc_t_min, air_calc_t_max = database.iloc[119, 1], database.iloc[119, 2], database.iloc[119, 3]
print(air_calc_p, air_calc_t_min, air_calc_t_max)
air_work_p, air_work_t = database.iloc[120, 1], database.iloc[120, 2]
print(air_work_p, air_work_t)


#-----------------------------------------------------------------------------------------------------------------------
# main
text = ['ООО «НТЦ «Ахмадуллины»',
        '']

for line in text:
    paragraph = doc.add_paragraph(line)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in paragraph.runs:
        set_font(run, 'Times New Roman', 18)
    set_paragraph_format(paragraph, left_indent=0.0, right_indent=0.0, first_line_indent=1.25, line_spacing=22,
                         space_after=0, space_before=0)

months_in_russian = {
    'January': 'Январь',
    'February': 'Февраль',
    'March': 'Март',
    'April': 'Апрель',
    'May': 'Май',
    'June': 'Июнь',
    'July': 'Июль',
    'August': 'Август',
    'September': 'Сентябрь',
    'October': 'Октябрь',
    'November': 'Ноябрь',
    'December': 'Декабрь'
}

current_date = datetime.datetime.now()

day = current_date.day
month_english = current_date.strftime("%B")
month = months_in_russian[month_english]
year = current_date.year

text = ['УТВЕРЖДАЮ',
        'Генеральный директор',
        '__________Р.М. Ахмадуллин',
       f'«{day}» {month} {year} года',
        '']

for line in text:
    paragraph = doc.add_paragraph(line)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    for run in paragraph.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph, left_indent=0.0, right_indent=0.0, first_line_indent=1.25, line_spacing=22,
                         space_after=0, space_before=0)

text = ['',
        'Базовый проект',
        'установки демеркаптанизации керосиновой фракции АО ««ННК-Хабаровский Нефтеперерабатывающий завод» по технологии «Demerus Jet»',
        '',
        '']

for line in text:
    paragraph = doc.add_paragraph(line)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in paragraph.runs:
        set_font(run, 'Times New Roman', 14, bold=False)
    set_paragraph_format(paragraph, left_indent=0.0, right_indent=0.0, first_line_indent=1.25, line_spacing=22,
                         space_after=0, space_before=0)

text = ['']

for line in text:
    paragraph = doc.add_paragraph(line)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in paragraph.runs:
        set_font(run, 'Times New Roman', 14, bold=True)
    set_paragraph_format(paragraph, left_indent=0.0, right_indent=0.0, first_line_indent=1.25, line_spacing=22,
                         space_after=0, space_before=0)

text = ['В настоящем документе содержится конфиденциальная информация относительно технологии «Demerus-Jet», включая эксплуатационные условия и технологические возможности, которые не могут быть раскрыты неуполномоченным лицам. Представленные материалы являются собственностью Лицензиара. Получая настоящую информацию, вы соглашаетесь не использовать ее ни для каких других целей, кроме тех, которые согласованы с Лицензиаром в письменной форме, не воспроизводить этот документ полностью или частично и не раскрывать его содержимое третьим лицам без письменного разрешения Лицензиара.',
        '',
        '']

# Добавление абзацев и установка их форматирования
for line in text:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = paragraph.add_run(line)
    set_font(run, 'Times New Roman', 12, italic=True)  # Установить курсив
    set_paragraph_format(paragraph, left_indent=0.0, right_indent=0.0, first_line_indent=1.25, line_spacing=18,
                         space_after=0, space_before=0)

table = doc.add_table(rows=6, cols=3)
table.style = 'Table Grid'

header_text = ['№ п/п', 'Ревизия', 'Дата выдачи']

hdr_cells = table.rows[0].cells
for i, text in enumerate(header_text):
    cell_paragraph = hdr_cells[i].paragraphs[0]
    cell_paragraph.text = text
    cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in cell_paragraph.runs:
        set_font(run, 'Times New Roman', 14, bold=True)
    set_paragraph_format(cell_paragraph, left_indent=0.0, right_indent=0.0, first_line_indent=0.0,
                         line_spacing=22, space_after=0, space_before=0)

text = ['',
        '',
        '',
        '',
        f'Казань – {year}']

for line in text:
    paragraph = doc.add_paragraph(line)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in paragraph.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph, left_indent=0.0, right_indent=0.0, first_line_indent=1.25, line_spacing=22,
                         space_after=0, space_before=0)

# Добавление нового раздела
new_section = doc.add_section(WD_SECTION.NEW_PAGE)

# Установка ориентации
new_section.orientation = WD_ORIENT.PORTRAIT

# Убедимся, что размеры страницы корректны для книжной ориентации
if new_section.page_width > new_section.page_height:
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width

ch_1 = head_counter.increment()

heading = doc.add_heading(f'{ch_1:.0f} ВВЕДЕНИЕ', level=1)
heading.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

for run in heading.runs:
    set_font(run, 'Times New Roman', 14)
    set_paragraph_format(heading, left_indent=0.0, right_indent=0.0, first_line_indent=1.25, line_spacing=22,
                         space_after=0, space_before=0)


text = ['',
        '',
        'Настоящий Базовый проект на проектирование установки очистки керосиновой фракции от меркаптанов и кислых примесей выполнен в соответствии с договором № 61 от 26 марта 2020г. для АО ««ННК-Хабаровский Нефтеперерабатывающий завод».',
        'Керосиновые фракции установок ЭЛОУ-АТ и ЭЛОУ-АВТ АО «ННК-Хабаровский Нефтеперерабатывающий завод» отличаются повышенным содержанием меркаптановой серы (от 300,0 до 400,0 ppm), не соответствующим требованиям ГОСТ 10227-86 «Топлива для реактивных двигателей» на авиационное топливо марки ТС-1 (не более 30,0 ppm по меркаптановой сере, сероводород - отсутствие).',
        'Меркаптановая сера в керосиновой фракции представлена высокомолекулярными соединениями, трудно извлекаемыми водно-щелочными растворами. Поэтому процесс их демеркаптанизации сводится к дезодорации содержащихся в них коррозионно-активных меркаптанов путем их окисления в инертные дисульфиды, остающиеся в очищаемом топливе. Такой подход оправдан приемлемым содержанием общей серы в этих фракциях.',
        f'Блок щелочной демеркаптанизации керосиновой фракции предназначен для окисления меркаптановых соединений и рассчитан на переработку по номинальной производительности {flow_rate} т/ч.',
        'В состав блока щелочной очистки «Demerus-Jet» входят:',
        '– узел окислительной демеркаптанизации керосиновой фракции;',
        '– узел адсорбционной очистки керосиновой фракции;',
        '– узел регенерации и концентрирования промотора;',
        '– реагентное хозяйство.',
        f'Режим работы блока демеркаптанизации керосиновой фракции «Demerus-Jet» - непрерывный, {work_time} часов в год. Расчетный период непрерывной эксплуатации установки между остановками на капитальный ремонт – 24 месяца. Срок службы оборудования не менее 20 лет. При расчете и подборе оборудования, согласно заданию на проектирование, был принят диапазон устойчивой производительности от {min_flow_rate} до {flow_rate} т/ч.',
        '']

for line in text:
    paragraph = doc.add_paragraph(line)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph, left_indent=0.0, right_indent=0.0, first_line_indent=1.25, line_spacing=22,
                         space_after=0, space_before=0)


#-----------------------------------------------------------------------------------------------------------------------
# Добавление нового раздела
new_section = doc.add_section(WD_SECTION.NEW_PAGE)

# Установка ориентации
new_section.orientation = WD_ORIENT.PORTRAIT

# Убедимся, что размеры страницы корректны для книжной ориентации
if new_section.page_width > new_section.page_height:
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width

ch_5 = head_counter.increment()

heading = doc.add_heading(f'{ch_5:.0f} ХАРАКТЕРИСТИКА ИСХОДНОГО СЫРЬЯ, ПРОДУКТОВ, ОСНОВНЫХ И ВСПОМОГАТЕЛЬНЫХ МАТЕРИАЛОВ', level=1)
heading.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
for run in heading.runs:
    set_font(run, 'Times New Roman', 14)
    set_paragraph_format(heading, left_indent=0.0, right_indent=0.0, first_line_indent=1.25, line_spacing=22,
                         space_after=0, space_before=0)

text = [f'',
        f'']

for line in text:
    paragraph_after_break = doc.add_paragraph(line)
    paragraph_after_break.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph_after_break.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph_after_break, left_indent=0.0, right_indent=0.0, first_line_indent=1.25,
                         line_spacing=22, space_after=0, space_before=0)

text = [f'Исходным сырьем блока демеркаптанизации керосиновой фракции "Demerus Jet" является прямогонный дистиллят (керосиновая фракция) в количестве от {min_flow_rate} до {flow_rate} т/ч и содержанием меркаптановой серы до {mass_frac_tiols}% мас. ({ppm_tiols} ppm).',
        f'']

for line in text:
    paragraph_after_break = doc.add_paragraph(line)
    paragraph_after_break.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph_after_break.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph_after_break, left_indent=0.0, right_indent=0.0, first_line_indent=1.25,
                         line_spacing=22, space_after=0, space_before=0)

table5_1 = table_counter.increment()
table5_2 = table_counter.increment()
table5_3 = table_counter.increment()
table5_4 = table_counter.increment()
table5_5 = table_counter.increment()

df5_1, merged_ranges = read_excel_with_merged_cells('database.xlsx', '5.1')

header_text_first = f'Таблица {table5_1:.1f} – Физико-химические показатели качества сырья, поступающего на блок "Demerus Jet"'
header_text_next = f'Продолжение таблицы {table5_1:.1f} – Физико-химические показатели качества сырья, поступающего на блок "Demerus Jet"'

rows_per_page_first = 100  # Количество строк для первой таблицы
rows_per_page_next = 100  # Количество строк для следующих таблиц

total_rows = len(df5_1)
start_row = 0

# Первая таблица с заголовком
end_row = min(start_row + rows_per_page_first, total_rows)
add_header(doc, header_text_first)
add_table(doc, df5_1, start_row, end_row, merged_ranges, include_header=True)
start_row = end_row

# Последующие таблицы без заголовка
while start_row < total_rows:
    end_row = min(start_row + rows_per_page_next, total_rows)
    insert_page_break(doc)
    add_header(doc, header_text_next)
    add_table(doc, df5_1, start_row, end_row, merged_ranges, include_header=True)
    start_row = end_row


text = [f'',
        f'',
        f'',
        f'',
        f'']

for line in text:
    paragraph_after_break = doc.add_paragraph(line)
    paragraph_after_break.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph_after_break.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph_after_break, left_indent=0.0, right_indent=0.0, first_line_indent=1.25,
                         line_spacing=22, space_after=0, space_before=0)

df5_2, merged_ranges = read_excel_with_merged_cells('database.xlsx', '5.2')

header_text_first = f'Таблица {table5_2:.1f} – Характеристика керосиновой фракции - сырья блока «Demerus Jet»'
header_text_next = f'Продолжение таблицы {table5_2:.1f} – Характеристика керосиновой фракции - сырья блока «Demerus Jet»'

rows_per_page_first = 100  # Количество строк для первой таблицы
rows_per_page_next = 100  # Количество строк для следующих таблиц

total_rows = len(df5_2)
start_row = 0

# Первая таблица с заголовком
end_row = min(start_row + rows_per_page_first, total_rows)
add_header(doc, header_text_first)
add_table(doc, df5_2, start_row, end_row, merged_ranges, include_header=True)
start_row = end_row

# Последующие таблицы без заголовка
while start_row < total_rows:
    end_row = min(start_row + rows_per_page_next, total_rows)
    insert_page_break(doc)
    add_header(doc, header_text_next)
    add_table(doc, df5_2, start_row, end_row, merged_ranges, include_header=True)
    start_row = end_row

text = [f'',
        f'Целевым продуктом блока "Demerus Jet" является керосиновая фракция с массовой долей меркаптановой серы не более 30 ppm, сероводород – отсутствие. Концентрация общей серы остается без изменений в диапазоне 0,114÷0,116 % мас.',
        f'',
        f'К основным материалам относятся (характеристики представлены в таблице {table5_5:.1f}):',
        f'- гетерогенный катализатор КСМ-Х, изготавливаемый в соответствии с ТУ 2175-001-40655797-2014',
        f'- глина отбеливающая (бентонитовая); ',
        f'- γ – оксид алюминия по ТУ 6-09-426-75;',
        f'- шары фарфоровые номинальный диаметр шара 3 мм, изготовляются в соответствии с ТУ 4328-030-07608911-2015. Материал - фарфор по ГОСТ 20419-83;',
        f'- воздух сжатый (КИП, технологический);',
        f'- пар среднего давления;',
        f'- пар низкого давления;',
        f'- оборотная вода прямая;',
        f'- оборотная вода обратная;',
        f'- инертный газ низкого давления (Азот);',
        f'- инертный газ высокого давления (Азот);',
        f'- деминерализованная вода для приготовления водных растворов NaOH и КОН;',
        f'- промотор КСП(ж), соответствует ТУ 0258-015-00151638-ОП-99. В качестве промотора КСП (тв.) используется калия гидрат окиси твердый – КОН. Промотор КСП(ж) образуется в ходе эксплуатации установки из продуктов взаимодействия кислых примесей керосина с гидроксидом калия и кислородом воздуха на поверхности гетерогенного катализатора КСМ-Х. Необходимость в закупки КСП(ж) отсутствует. ',
        f'Промотор КСП(ж) представляет собой темно-коричневую жидкость с плотностью не менее 1,3 кг/дм3. При гравиметрическом отстаивании он расслаивается на два слоя: светлый тяжелый (КСП(ж)) и темный легкий (калиевые соли нафтеновых кислот). Хранится при температуре не ниже 5оС. ',
        f'В состав промотора КСП(ж) входит спектр органических кислых примесей, извлеченных щелочью из керосиновой фракции и окисленных на гетерогенном катализаторе КСМ-Х воздухом до алкилтиосульфонатов, солей сульфокислот и др. кислородсодержащих продуктов.',
        f'']

for line in text:
    paragraph_after_break = doc.add_paragraph(line)
    paragraph_after_break.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph_after_break.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph_after_break, left_indent=0.0, right_indent=0.0, first_line_indent=1.25,
                         line_spacing=22, space_after=0, space_before=0)

df5_3, merged_ranges = read_excel_with_merged_cells('database.xlsx', '5.3')

header_text_first = f'Таблица {table5_3:.1f} – Характеристика керосиновой фракции - сырья блока «Demerus Jet»'
header_text_next = f'Продолжение таблицы {table5_3:.1f} – Характеристика керосиновой фракции - сырья блока «Demerus Jet»'

rows_per_page_first = 100  # Количество строк для первой таблицы
rows_per_page_next = 100  # Количество строк для следующих таблиц

total_rows = len(df5_3)
start_row = 0

# Первая таблица с заголовком
end_row = min(start_row + rows_per_page_first, total_rows)
add_header(doc, header_text_first)
add_table(doc, df5_3, start_row, end_row, merged_ranges, include_header=True)
start_row = end_row

# Последующие таблицы без заголовка
while start_row < total_rows:
    end_row = min(start_row + rows_per_page_next, total_rows)
    insert_page_break(doc)
    add_header(doc, header_text_next)
    add_table(doc, df5_3, start_row, end_row, merged_ranges, include_header=True)
    start_row = end_row

text = [f'']

for line in text:
    paragraph_after_break = doc.add_paragraph(line)
    paragraph_after_break.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph_after_break.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph_after_break, left_indent=0.0, right_indent=0.0, first_line_indent=1.25,
                         line_spacing=22, space_after=0, space_before=0)

df5_4, merged_ranges = read_excel_with_merged_cells('database.xlsx', '5.4')

header_text_first = f'Таблица {table5_4:.1f} – Физико-химические характеристики КСП (ж)'
header_text_next = f'Продолжение таблицы {table5_4:.1f} – Физико-химические характеристики КСП (ж)'

rows_per_page_first = 100  # Количество строк для первой таблицы
rows_per_page_next = 100  # Количество строк для следующих таблиц

total_rows = len(df5_4)
start_row = 0

# Первая таблица с заголовком
end_row = min(start_row + rows_per_page_first, total_rows)
add_header(doc, header_text_first)
add_table(doc, df5_4, start_row, end_row, merged_ranges, include_header=True)
start_row = end_row

# Последующие таблицы без заголовка
while start_row < total_rows:
    end_row = min(start_row + rows_per_page_next, total_rows)
    insert_page_break(doc)
    add_header(doc, header_text_next)
    add_table(doc, df5_4, start_row, end_row, merged_ranges, include_header=True)
    start_row = end_row

text = [f''
       ]

for line in text:
    paragraph_after_break = doc.add_paragraph(line)
    paragraph_after_break.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph_after_break.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph_after_break, left_indent=0.0, right_indent=0.0, first_line_indent=1.25,
                         line_spacing=22, space_after=0, space_before=0)


# Добавление нового раздела
new_section = doc.add_section(WD_SECTION.NEW_PAGE)

# Установка ориентации
new_section.orientation = WD_ORIENT.PORTRAIT

# Убедимся, что размеры страницы корректны для альбомной ориентации
if new_section.page_width < new_section.page_height:
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width

df5_5, merged_ranges = read_excel_with_merged_cells('database.xlsx', '5.5')

header_text_first = f'Таблица {table5_5:.1f} – Характеристика основных и вспомогательных материалов'
header_text_next = f'Продолжение таблицы {table5_5:.1f} – Характеристика основных и вспомогательных материалов'

rows_per_page_first = 100  # Количество строк для первой таблицы
rows_per_page_next = 5  # Количество строк для следующих таблиц

total_rows = len(df5_5)
start_row = 0

# Первая таблица с заголовком
end_row = min(start_row + rows_per_page_first, total_rows)
add_header(doc, header_text_first)
add_table(doc, df5_5, start_row, end_row, merged_ranges, include_header=False)
start_row = end_row

# Последующие таблицы без заголовка
while start_row < total_rows:
    end_row = min(start_row + rows_per_page_next, total_rows)
    insert_page_break(doc)
    add_header(doc, header_text_next)
    add_table(doc, df5_5, start_row, end_row, merged_ranges, include_header=False)
    start_row = end_row

text = [f''
       ]

for line in text:
    paragraph_after_break = doc.add_paragraph(line)
    paragraph_after_break.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph_after_break.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph_after_break, left_indent=0.0, right_indent=0.0, first_line_indent=1.25,
                         line_spacing=22, space_after=0, space_before=0)


# Добавление нового раздела
new_section = doc.add_section(WD_SECTION.NEW_PAGE)

# Установка ориентации
new_section.orientation = WD_ORIENT.PORTRAIT

# Убедимся, что размеры страницы корректны для альбомной ориентации
if new_section.page_width > new_section.page_height:
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width

ch_6 = head_counter.increment()

heading = doc.add_heading(f'{ch_6:.0f} ТЕХНИЧЕСКАЯ ХАРАКТЕРИСТИКА ОТХОДОВ И ОТРАБОТАННОГО ВОЗДУХА', level=1)
heading.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
for run in heading.runs:
    set_font(run, 'Times New Roman', 14)
    set_paragraph_format(heading, left_indent=0.0, right_indent=0.0, first_line_indent=1.25, line_spacing=22,
                         space_after=0, space_before=0)

text = [f'',
        f'']

for line in text:
    paragraph_after_break = doc.add_paragraph(line)
    paragraph_after_break.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph_after_break.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph_after_break, left_indent=0.0, right_indent=0.0, first_line_indent=1.25,
                         line_spacing=22, space_after=0, space_before=0)


text = [f'Отходами или выбросами с блока «Demerus Jet» могут быть:',
        '1) Отработанная глина. Срок эксплуатации – 6 месяцев. Вывозят специализированными организациями на утилизацию.',
        '2) Отработанный γ-оксид алюминия. Срок эксплуатации – 6 месяцев. Вывозят специализированными организациями на утилизацию. ',
        '3) Отработанный катализатор КСМ-Х, при потере его стабильности вследствие нарушения предписанных ТУ правил хранения и эксплуатации, либо по истечении гарантийного срока службы катализатора. Код отхода по Федеральному классификационному каталогу отходов ФККО 2017 - 4 41 006 01 49 3 катализатор на основе пропилена с содержанием фталоциандисульфата кобальта менее 15,0 % отработанный.',
        '4) Отработанные фарфоровые шары. Срок эксплуатации – 8 лет. Направляются на полигон для захоронения.',
        ''
       ]

for line in text:
    paragraph_after_break = doc.add_paragraph(line)
    paragraph_after_break.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph_after_break.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph_after_break, left_indent=0.0, right_indent=0.0, first_line_indent=1.25,
                         line_spacing=22, space_after=0, space_before=0)

table6_1 = table_counter.increment()
table6_2 = table_counter.increment()
table6_3 = table_counter.increment()

df6_1, merged_ranges = read_excel_with_merged_cells('database.xlsx', '6.1')

header_text_first = f'Таблица {table6_1:.1f} – Характеристика побочных продуктов и выбросов в пересчёте на тонну перерабатываемого сырья'
header_text_next = f'Продолжение таблицы {table6_1:.1f} – Характеристика побочных продуктов и выбросов в пересчёте на тонну перерабатываемого сырья'

rows_per_page_first = 100  # Количество строк для первой таблицы
rows_per_page_next = 100  # Количество строк для следующих таблиц

total_rows = len(df6_1)
start_row = 0

# Первая таблица с заголовком
end_row = min(start_row + rows_per_page_first, total_rows)
add_header(doc, header_text_first)
add_table(doc, df6_1, start_row, end_row, merged_ranges, include_header=True)
start_row = end_row

# Последующие таблицы без заголовка
while start_row < total_rows:
    end_row = min(start_row + rows_per_page_next, total_rows)
    insert_page_break(doc)
    add_header(doc, header_text_next)
    add_table(doc, df6_1, start_row, end_row, merged_ranges, include_header=True)
    start_row = end_row

text = [f''
       ]

for line in text:
    paragraph_after_break = doc.add_paragraph(line)
    paragraph_after_break.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph_after_break.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph_after_break, left_indent=0.0, right_indent=0.0, first_line_indent=1.25,
                         line_spacing=22, space_after=0, space_before=0)

df6_2, merged_ranges = read_excel_with_merged_cells('database.xlsx', '6.2')

header_text_first = f'Таблица {table6_2:.1f} – Техническая характеристика побочных продуктов'
header_text_next = f'Продолжение таблицы {table6_2:.1f} – Техническая характеристика побочных продуктов'

rows_per_page_first = 100  # Количество строк для первой таблицы
rows_per_page_next = 100  # Количество строк для следующих таблиц

total_rows = len(df6_2)
start_row = 0

# Первая таблица с заголовком
end_row = min(start_row + rows_per_page_first, total_rows)
add_header(doc, header_text_first)
add_table(doc, df6_2, start_row, end_row, merged_ranges, include_header=True)
start_row = end_row

# Последующие таблицы без заголовка
while start_row < total_rows:
    end_row = min(start_row + rows_per_page_next, total_rows)
    insert_page_break(doc)
    add_header(doc, header_text_next)
    add_table(doc, df6_2, start_row, end_row, merged_ranges, include_header=True)
    start_row = end_row

text = [f''
       ]

for line in text:
    paragraph_after_break = doc.add_paragraph(line)
    paragraph_after_break.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph_after_break.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph_after_break, left_indent=0.0, right_indent=0.0, first_line_indent=1.25,
                         line_spacing=22, space_after=0, space_before=0)

df6_3, merged_ranges = read_excel_with_merged_cells('database.xlsx', '6.3')

header_text_first = f'Таблица {table6_3:.1f} – Условия сбора, хранения, транспортирования, складирования и захоронения отходов'
header_text_next = f'Продолжение таблицы {table6_3:.1f} – Условия сбора, хранения, транспортирования, складирования и захоронения отходов'

rows_per_page_first = 100  # Количество строк для первой таблицы
rows_per_page_next = 100  # Количество строк для следующих таблиц

total_rows = len(df6_3)
start_row = 0

# Первая таблица с заголовком
end_row = min(start_row + rows_per_page_first, total_rows)
add_header(doc, header_text_first)
add_table(doc, df6_3, start_row, end_row, merged_ranges, include_header=True)
start_row = end_row

# Последующие таблицы без заголовка
while start_row < total_rows:
    end_row = min(start_row + rows_per_page_next, total_rows)
    insert_page_break(doc)
    add_header(doc, header_text_next)
    add_table(doc, df6_3, start_row, end_row, merged_ranges, include_header=True)
    start_row = end_row

text = [f''
       ]

for line in text:
    paragraph_after_break = doc.add_paragraph(line)
    paragraph_after_break.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph_after_break.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph_after_break, left_indent=0.0, right_indent=0.0, first_line_indent=1.25,
                         line_spacing=22, space_after=0, space_before=0)


# Добавление нового раздела
new_section = doc.add_section(WD_SECTION.NEW_PAGE)

# Установка ориентации
new_section.orientation = WD_ORIENT.PORTRAIT

# Убедимся, что размеры страницы корректны для книжной ориентации
if new_section.page_width > new_section.page_height:
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width

ch_10 = head_counter.increment()

heading = doc.add_heading(f'{ch_10:.0f} МАТЕРИАЛЬНЫЙ БАЛАНС ПРОЦЕССА', level=1)
heading.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
for run in heading.runs:
    set_font(run, 'Times New Roman', 14)
    set_paragraph_format(heading, left_indent=0.0, right_indent=0.0, first_line_indent=1.25, line_spacing=22,
                         space_after=0, space_before=0)

text = [f'',
        f'']

for line in text:
    paragraph_after_break = doc.add_paragraph(line)
    paragraph_after_break.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph_after_break.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph_after_break, left_indent=0.0, right_indent=0.0, first_line_indent=1.25,
                         line_spacing=22, space_after=0, space_before=0)

text = [f'Исходные данные для расчета материального баланса',
        f'',
        f'Материальный баланс установки демеркаптанизации керосиновой фракции («Demerus-Jet») составлен в соответствии со следующим расчетом:',
        f'']

for line in text:
    paragraph_after_break = doc.add_paragraph(line)
    paragraph_after_break.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph_after_break.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph_after_break, left_indent=0.0, right_indent=0.0, first_line_indent=1.25,
                         line_spacing=22, space_after=0, space_before=0)

# Добавление нового раздела
new_section = doc.add_section(WD_SECTION.NEW_PAGE)

# Установка ориентации
new_section.orientation = WD_ORIENT.PORTRAIT

# Убедимся, что размеры страницы корректны для альбомной ориентации
if new_section.page_width < new_section.page_height:
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width

table10_1 = table_counter.increment()
table10_2 = table_counter.increment()
table10_3 = table_counter.increment()

df10_1, merged_ranges = read_excel_with_merged_cells('database.xlsx', '10.1')

header_text_first = f'Таблица {table10_1:.1f} – Материальный баланс установки демеркаптанизации керосиновой фракции'
header_text_next = f'Продолжение таблицы {table10_1:.1f} – Материальный баланс установки демеркаптанизации керосиновой фракции'

rows_per_page_first = 18  # Количество строк для первой таблицы
rows_per_page_next = 18  # Количество строк для следующих таблиц

total_rows = len(df10_1)
start_row = 0

# Первая таблица с заголовком
end_row = min(start_row + rows_per_page_first, total_rows)
add_header(doc, header_text_first)
add_table(doc, df10_1, start_row, end_row, merged_ranges, include_header=False)
start_row = end_row

# Последующие таблицы без заголовка
while start_row < total_rows:
    end_row = min(start_row + rows_per_page_next, total_rows)
    insert_page_break(doc)
    add_header(doc, header_text_next)
    add_table(doc, df10_1, start_row, end_row, merged_ranges, include_header=False)
    start_row = end_row

#-----------------------------------------------------------------------------------------------------------------------

# Сохраняем документ
doc.save('БП.docx')
