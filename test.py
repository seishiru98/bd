from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION, WD_ORIENT

import pandas as pd
from openpyxl import load_workbook

import datetime
#-----------------------------------------------------------------------------------------------------------------------

doc = Document()

section = doc.sections[0]

section.left_margin = Cm(2)  # Левое поле
section.right_margin = Cm(1)  # Правое поле
section.top_margin = Cm(2)  # Верхнее поле
section.bottom_margin = Cm(2)  # Нижнее поле
#-----------------------------------------------------------------------------------------------------------------------
# Функции

def set_font(run, font_name, font_size, italic=False, bold=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.italic = italic
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    # Это нужно для корректного отображения шрифта на всех платформах
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.append(rFonts)


def set_paragraph_format(paragraph, left_indent=0, right_indent=0, first_line_indent=1.25, line_spacing=22,
                         space_after=0, space_before=0):
    paragraph_format = paragraph.paragraph_format
    paragraph_format.left_indent = Cm(left_indent)
    paragraph_format.right_indent = Cm(right_indent)
    paragraph_format.first_line_indent = Cm(first_line_indent)
    paragraph_format.line_spacing = Pt(line_spacing)
    paragraph_format.space_after = Cm(space_after)
    paragraph_format.space_before = Cm(space_before)


def add_header(doc, header_text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(header_text)
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY


from docx.shared import Cm


def add_table(doc, df, start_row, end_row, merged_ranges, include_header=True):
    total_width = Cm(18.5)  # Общая ширина таблицы (например, вся ширина страницы)
    max_col_widths = []

    # Определим максимальную длину текста для каждой колонки
    for col in df.columns:
        max_width = max(df[col][start_row:end_row].apply(lambda x: len(str(x)) if x is not None else 0))
        max_col_widths.append(max_width)

    # Пропорциональная ширина колонок относительно их максимального содержания
    total_content_width = sum(max_col_widths)
    col_width_ratios = [width / total_content_width for width in max_col_widths]
    col_widths = [total_width * ratio for ratio in col_width_ratios]

    # Создаем таблицу с количеством столбцов, равным числу колонок в DataFrame
    num_columns = len(df.columns)
    table = doc.add_table(rows=0, cols=num_columns)
    table.style = 'Table Grid'

    # Добавляем заголовок таблицы, если включен параметр include_header
    if include_header:
        hdr_cells = table.add_row().cells
        for i, column_name in enumerate(df.columns):
            hdr_cells[i].width = col_widths[i]  # Применяем пропорциональную ширину к заголовкам
            cell_paragraph = hdr_cells[i].paragraphs[0]
            cell_paragraph.text = str(column_name)
            cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            for run in cell_paragraph.runs:
                set_font(run, 'Times New Roman', 12, bold=False)
            set_paragraph_format(cell_paragraph, left_indent=0.0, right_indent=0.0,
                                 first_line_indent=0.0, line_spacing=18, space_after=0, space_before=0)

    # Добавление строк таблицы
    for index in range(start_row, end_row):
        row = df.iloc[index]
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].width = col_widths[i]  # Применяем пропорциональную ширину к каждой ячейке строки
            cell_paragraph = row_cells[i].paragraphs[0]
            cell_paragraph.text = str(value)
            cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            for run in cell_paragraph.runs:
                set_font(run, 'Times New Roman', 12)
            set_paragraph_format(cell_paragraph, left_indent=0.0, right_indent=0.0,
                                 first_line_indent=0.0, line_spacing=18, space_after=0, space_before=0)

    # Корректировка для индексации строк
    header_offset = 1 if include_header else 0

    # Обработка объединённых ячеек
    for merged_range in merged_ranges:
        min_col, min_row, max_col, max_row = merged_range.bounds
        min_col -= 1  # Преобразование к нулевой базе
        max_col -= 1

        # Корректировка индексов строк для DataFrame
        min_row_df = min_row - 2  # -2, потому что DataFrame начинается с Excel строки 2 и индексируется с 0
        max_row_df = max_row - 2

        # Проверяем, попадает ли объединённый диапазон в текущий диапазон строк
        if min_row_df >= start_row and max_row_df < end_row:
            start_cell_row = (min_row_df - start_row) + header_offset
            end_cell_row = (max_row_df - start_row) + header_offset

            start_cell_col = min_col
            end_cell_col = max_col

            start_cell = table.cell(int(start_cell_row), int(start_cell_col))
            end_cell = table.cell(int(end_cell_row), int(end_cell_col))
            start_cell.merge(end_cell)

            # Удаляем лишние пустые параграфы из объединенной ячейки
            for paragraph in start_cell.paragraphs:
                if not paragraph.text.strip():
                    p = paragraph._element
                    p.getparent().remove(p)
                    p._p = p._element = None


def insert_page_break(doc):
    doc.add_page_break()


def read_excel_with_merged_cells(filename, sheet_name):
    # Загружаем Excel-файл с помощью openpyxl
    wb = load_workbook(filename, data_only=True)
    ws = wb[sheet_name]

    # Сохраняем данные в список строк (list of lists)
    data = []
    for row in ws.iter_rows(values_only=True):
        data.append(list(row))

    # Получаем объединенные ячейки
    merged_ranges = ws.merged_cells.ranges

    # Проходим по каждому объединенному диапазону и заполняем его данные
    for merged_range in merged_ranges:
        # Получаем диапазон объединённых ячеек
        min_col, min_row, max_col, max_row = merged_range.bounds

        # Получаем значение из первой ячейки диапазона
        merged_value = ws.cell(row=min_row, column=min_col).value

        # Присваиваем это значение только первой ячейке, остальные оставляем пустыми
        for row in range(min_row - 1, max_row):
            for col in range(min_col - 1, max_col):
                if row == min_row - 1 and col == min_col - 1:
                    data[row][col] = merged_value
                else:
                    data[row][col] = ''  # Оставляем остальные ячейки пустыми

    # Преобразуем данные в DataFrame pandas
    df = pd.DataFrame(data[1:], columns=data[0])
    return df, merged_ranges


class Counter:
    def __init__(self, start_value, step):
        self.value = start_value
        self.step = step

    def increment(self):
        current_value = self.value
        self.value += self.step
        return current_value

class HeadingCounter(Counter):
    def __init__(self, start_value, paragraph_counter, table_counter, fig_counter):
        super().__init__(start_value, 1)
        self.paragraph_counter = paragraph_counter
        self.table_counter = table_counter
        self.fig_counter = fig_counter

    def increment(self):
        current_value = super().increment()
        self.paragraph_counter.reset(current_value + 0.1)
        self.table_counter.reset(current_value + 0.1)
        self.fig_counter.reset(current_value + 0.1)
        return current_value

class ParagraphCounter(Counter):
    def reset(self, new_start_value):
        self.value = new_start_value

class TableCounter(Counter):
    def reset(self, new_start_value):
        self.value = new_start_value

class FigCounter(Counter):
    def reset(self, new_start_value):
        self.value = new_start_value

# Инициализация счетчиков
n_heading = 1
n_paragraph_start = n_heading + 0.1
n_table_start = n_heading + 0.1
n_fig_start = n_heading + 0.1

par_counter = ParagraphCounter(n_paragraph_start, 0.1)
table_counter = TableCounter(n_table_start, 0.1)
fig_counter = FigCounter(n_fig_start, 0.1)
head_counter = HeadingCounter(n_heading, par_counter, table_counter, fig_counter)

def read_excel_data(filename, sheet_name):
    try:
        return pd.read_excel(filename, sheet_name=sheet_name)
    except FileNotFoundError:
        print(f"ОШИБКА: Файл {filename} не найден.")
    except ValueError:
        print(f"ОШИБКА: Лист {sheet_name} не найден в файле {filename}.")
    except Exception as e:
        print(f"ОШИБКА: Произошла ошибка при чтении файла {filename}: {e}")

#-----------------------------------------------------------------------------------------------------------------------

ch_10 = head_counter.increment()

heading = doc.add_heading(f'{ch_10:.0f} МАТЕРИАЛЬНЫЙ БАЛАНС ПРОЦЕССА', level=1)
heading.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
for run in heading.runs:
    set_font(run, 'Times New Roman', 14)
    set_paragraph_format(heading, left_indent=0.0, right_indent=0.0, first_line_indent=1.25, line_spacing=22,
                         space_after=0, space_before=0)

text = [f'',
        f'']

for line in text:
    paragraph_after_break = doc.add_paragraph(line)
    paragraph_after_break.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph_after_break.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph_after_break, left_indent=0.0, right_indent=0.0, first_line_indent=1.25,
                         line_spacing=22, space_after=0, space_before=0)

text = [f'Исходные данные для расчета материального баланса',
        f'',
        f'Материальный баланс установки демеркаптанизации керосиновой фракции («Demerus-Jet») составлен в соответствии со следующим расчетом:',
        f'']

for line in text:
    paragraph_after_break = doc.add_paragraph(line)
    paragraph_after_break.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    for run in paragraph_after_break.runs:
        set_font(run, 'Times New Roman', 14)
    set_paragraph_format(paragraph_after_break, left_indent=0.0, right_indent=0.0, first_line_indent=1.25,
                         line_spacing=22, space_after=0, space_before=0)

# Добавление нового раздела
new_section = doc.add_section(WD_SECTION.NEW_PAGE)

# Установка ориентации
new_section.orientation = WD_ORIENT.PORTRAIT

# Убедимся, что размеры страницы корректны для альбомной ориентации
if new_section.page_width < new_section.page_height:
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width


# Чтение данных из Excel с учётом объединённых ячеек
df2_1, merged_ranges = read_excel_with_merged_cells('database.xlsx', '10.1')

# Инициализация счётчика для таблиц
ch_2_par_1 = table_counter.increment()

header_text_first = f'Таблица {ch_2_par_1:.1f} – Содержание общей серы в СУГ после демеркаптанизации на гомогенных и гетерогенных катализаторах'
header_text_next = f'Продолжение таблицы {ch_2_par_1:.1f} – Содержание общей серы в СУГ после демеркаптанизации на гомогенных и гетерогенных катализаторах'

rows_per_page_first = 18  # Количество строк для первой таблицы
rows_per_page_next = 18  # Количество строк для следующих таблиц

total_rows = len(df2_1)
start_row = 0

# Первая таблица с заголовком
end_row = min(start_row + rows_per_page_first, total_rows)
add_header(doc, header_text_first)
add_table(doc, df2_1, start_row, end_row, merged_ranges, include_header=False)
start_row = end_row

# Последующие таблицы без заголовка
while start_row < total_rows:
    end_row = min(start_row + rows_per_page_next, total_rows)
    insert_page_break(doc)
    add_header(doc, header_text_next)
    add_table(doc, df2_1, start_row, end_row, merged_ranges, include_header=False)
    start_row = end_row

# Добавление нового раздела
new_section = doc.add_section(WD_SECTION.NEW_PAGE)

# Установка ориентации
new_section.orientation = WD_ORIENT.PORTRAIT

# Убедимся, что размеры страницы корректны для альбомной ориентации
if new_section.page_width > new_section.page_height:
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width


# Чтение данных из Excel с учётом объединённых ячеек
df2_1, merged_ranges = read_excel_with_merged_cells('database.xlsx', '5.3')

# Инициализация счётчика для таблиц
ch_2_par_1 = table_counter.increment()

header_text_first = f'Таблица {ch_2_par_1:.1f} – Содержание общей серы в СУГ после демеркаптанизации на гомогенных и гетерогенных катализаторах'
header_text_next = f'Продолжение таблицы {ch_2_par_1:.1f} – Содержание общей серы в СУГ после демеркаптанизации на гомогенных и гетерогенных катализаторах'

rows_per_page_first = 100  # Количество строк для первой таблицы
rows_per_page_next = 18  # Количество строк для следующих таблиц

total_rows = len(df2_1)
start_row = 0

# Первая таблица с заголовком
end_row = min(start_row + rows_per_page_first, total_rows)
add_header(doc, header_text_first)
add_table(doc, df2_1, start_row, end_row, merged_ranges, include_header=True)
start_row = end_row

# Последующие таблицы без заголовка
while start_row < total_rows:
    end_row = min(start_row + rows_per_page_next, total_rows)
    insert_page_break(doc)
    add_header(doc, header_text_next)
    add_table(doc, df2_1, start_row, end_row, merged_ranges, include_header=False)
    start_row = end_row
#-----------------------------------------------------------------------------------------------------------------------

# Сохраняем документ
doc.save('test.docx')
